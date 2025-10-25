"""DOCX exporter."""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document as DocxDocument
from docx.shared import Pt

from ..schemas.document import Document, DocumentBlock, DocumentPage
from .base import DocumentExporter, ExportFormat, ExportRequest, ExportResult


@dataclass
class DocxExporter(DocumentExporter):
    format: ExportFormat = ExportFormat.DOCX

    def export(self, document: Document, request: ExportRequest) -> ExportResult:
        doc = DocxDocument()
        doc.core_properties.comments = "Exported by Reflow OCR"
        for page in document.pages:
            self._render_page(doc, page)
        buffer = io.BytesIO()
        doc.save(buffer)
        filename = f"{request.filename_hint}.docx"
        return ExportResult(filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", content=buffer.getvalue())

    def _render_page(self, doc: DocxDocument, page: DocumentPage) -> None:
        doc.add_heading(f"Страница {page.index + 1}", level=2)
        for block in page.blocks:
            self._render_block(doc, block)
        doc.add_page_break()

    def _render_block(self, doc: DocxDocument, block: DocumentBlock) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        text = "\n".join(span.text for span in block.spans if span.text)
        if not text:
            text = "[пустой блок]"
        run = paragraph.add_run(text)
        if block.type == "header":
            run.bold = True
